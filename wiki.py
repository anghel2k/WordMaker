import wikipedia
import re
from docx import Document

title = input("Despre ce vrei sa fie proiectul?\n")
name =  input("Introdu numele tau: ")
wikipedia.set_lang("ro")
wiki = wikipedia.page('Circuit Integrat')
while True:
	try:
		wiki = wikipedia.page(title)
		break
	except:
		print("Nume proiect invalid")

text = wiki.content
text = re.sub(r'==', '', text)
text = re.sub(r'\n', '\n    ', text)
split = text.split('Vezi și', 1)
text = split[0]
# print(text)

document = Document()
paragraph = document.add_heading(title, 0)
paragraph.alignment = 1

document.add_paragraph('     ' + text)
paragraph = document.add_paragraph(name)
paragraph.alignment = 2
document.save(title + ".docx")